# -*- coding: utf-8 -*-
"""
    sphinxpapyrus.docxwriter.writer
    ~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~

    Custom docutils writer for docx.

    :copyright: Copyright 2018 by nakandev.
    :license: MIT, see LICENSE for details.
"""

import sys
import os
import re

from docutils import nodes, writers

from sphinx import addnodes
from sphinx.locale import admonitionlabels, _
from sphinx.util import logging

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER

package_dir = os.path.abspath(os.path.dirname(__file__))

logger = logging.getLogger(__name__)

# monkey patch
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne
from docx.oxml.simpletypes import XsdInt
from docx.oxml import register_element_cls
class CT_TrPr(BaseOxmlElement):
    tblHeader = ZeroOrOne('w:tblHeader')
register_element_cls('w:trPr', CT_TrPr)

class DocxWriter(writers.Writer):
    supported = ('docx',)
    settings_spec = ('No options here.', '', ())
    settings_defaults = {}  # type: Dict

    output = None

    def __init__(self, builder):
        # type: (DocxBuilder) -> None
        writers.Writer.__init__(self)
        self.builder = builder

        stylefile = builder.config.docx_style
        if stylefile:
            style_dir = self.builder.srcdir
            style_fullpath = os.path.join(style_dir, stylefile)
            self.docx = Document(style_fullpath)
        else:
            style_dir = os.path.join(package_dir, 'templates')
            style_fullpath = os.path.join(style_dir, 'style.docx')
            self.docx = Document(style_fullpath)
        self.docx_set_coreproperties()
        self.docx._body.clear_content()

    def docx_set_coreproperties(self):
        new_coreprop = self.builder.config.docx_coreproperties
        for name, value in new_coreprop.items():
            setattr(self.docx.core_properties, name, value)

    def translate(self):
        # type: () -> None
        visitor = self.builder.create_translator(self.document, self.builder, self.docx)
        self.document.walkabout(visitor)
        self.output = visitor.body

    def save(self, filename):
        self.docx.save(filename)

class DocxTranslator(nodes.NodeVisitor):

    stylename = {
        # paragraph styles
        'title': 'Title',
        'subtitle': 'Subtitle',
        'heading_': ['Heading %d' % (i+1) for i in range(6)],
        'bullet_list': 'Bullet List',
        'enumerated_list': 'Enumerated List',
        '_empty_bullet_list': 'Empty Bullet List',
        'definition_list': 'Definition List',
        'definition_list_item': 'Definition List Item',
        'field_list': 'Field List',
        'field_list_item': 'Field List Item',
        'option_list': 'Option List',
        'option_list_item': 'Option List Item',
        'literal_block': 'Literal Block',
        'block_quote': 'Quote',
        'line_block': 'Line Block',
        'doctest_block': 'Doctest Block',
        'transition': 'Horizontal Line',
        'table_caption': 'Caption',
        'image_caption': 'Caption',
        'code_block_caption': 'Code Block Caption',
        # character styles
        'strong': 'Strong',
        'emphasis': 'Emphasis',
        'literal_emphasis': 'Literal Emphasis',
        'subscript': 'Subscript',
        'superscript': 'Superscript',
        'title_reference': 'Book Title',
        'literal': 'Literal',
        'reference': 'Hyperlink',
        'footnote_reference': 'Default Paragraph Font',
        # table styles
        'table': ['Sphinx Table Normal', 'Sphinx Table List'],
    }

    def __init__(self, document, builder, docx):
        # type: (nodes.Node, DocxBuilder) -> None
        nodes.NodeVisitor.__init__(self, document)
        self.builder = builder
        self.settings = document.settings
        self.docnames = [builder.current_docname]
        self.docx = docx
        self.numbered = 0
        self.numbered_level = 0
        self.section_level = 0
        self.section_numIds = [self._get_new_num(abstractNumId=12)]
        self.initial_header_level = 0  # int(self.settings.initial_header_level)
        # docx paragraph properties
        self.p = None
        self.p_parents = [self.docx]
        self.p_style = []
        self.p_level = 0
        self.numIds = []
        self.is_first_list_item = False
        # special paragraphs
        self.tables = []
        self.item_width_rate = 0.8
        # docx run properties
        self.r = None
        self.r_style = None

    def _fignum_prefix(self, node):
        prefix = ''
        if self.builder.config.numfig:
            figtype = self.builder.env.domains['std'].get_figtype(node)
            format = self.builder.config.numfig_format.get(figtype)
            alias = u'%s/%s' % (self.docnames[-1], figtype)
            id = node['ids'][0]
            nums = self.builder.fignumbers[alias][id]
            prefix = format % '.'.join(map(str, nums))
        return prefix

    def _add_paragraph(self, text=None, style=None):
        p = None
        try:
            if isinstance(style, list):
                p = self.p_parents[-1].add_paragraph(text, style[-1])
            else:
                p = self.p_parents[-1].add_paragraph(text, style)
        except:
            p = self.p_parents[-1].add_paragraph(text, 'Normal')
        if self.p_level > 0:
            self._multilevel_list_numbering(p, self.p_level - 1, 15)
        return p

    def _add_run(self, text=None, style=None):
        r = None
        if self.p:
            try:
                r = self.p.add_run(text, style)
            except:
                r = self.p.add_run(text, 'Default Paragraph Font')
        return r

    def _get_new_num(self, abstractNumId):
        # monkey patch
        from types import MethodType
        from docx.parts.numbering import _NumberingDefinitions
        from docx.oxml.numbering import CT_Num
        from docx.oxml.numbering import CT_Numbering
        def add_num(self, abstractNum_id, restart=False):
            next_num_id = self._next_numId
            num = CT_Num.new(next_num_id, abstractNum_id)
            if restart:
                num.add_lvlOverride(ilvl=0).add_startOverride(1)
            return self._insert_num(num)
        numbering = self.docx._part.numbering_part.numbering_definitions._numbering
        if sys.version_info.major >= 3:
            numbering.add_num = MethodType(add_num, numbering)
        else:
            numbering.add_num = MethodType(add_num, numbering, CT_Numbering)
        num = numbering.add_num(abstractNumId, True).numId
        return num

    def _multilevel_list_numbering(self, paragraph, ilvl, numId):
        # monkey patch
        pfmt = paragraph.paragraph_format
        numPr = pfmt._element.get_or_add_pPr().get_or_add_numPr()
        numPr.get_or_add_ilvl().val = ilvl
        numPr.get_or_add_numId().val = numId

    def visit_start_of_file(self, node):
        # type: (nodes.Node) -> None
        self.docnames.append(node['docname'])

    def depart_start_of_file(self, node):
        # type: (nodes.Node) -> None
        self.docnames.pop()

    def visit_document(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_document(self, node):
        # type: (nodes.Node) -> None
        self.body = 'dommy text'

    def visit_highlightlang(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_section(self, node):
        # type: (nodes.Node) -> None
        self.section_level += 1
        if self.numbered:
            numId = self._get_new_num(abstractNumId=12)
            self.section_numIds.append(numId)

    def depart_section(self, node):
        # type: (nodes.Node) -> None
        self.section_level -= 1
        if self.numbered:
            self.section_numIds.pop()

    def visit_topic(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_topic(self, node):
        # type: (nodes.Node) -> None
        pass

    #visit_sidebar = visit_topic
    #depart_sidebar = depart_topic

    def visit_rubric(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_rubric(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_compound(self, node):
        # type: (nodes.Node) -> None
        if 'toctree-wrapper' in node.get('classes', None):
            numbered = node.get('numbered', 0)
            if numbered > 0 and self.numbered_level <= self.section_level:
                self.numbered = numbered
                self.numbered_level = self.section_level

    def depart_compound(self, node):
        # type: (nodes.Node) -> None
        if 'toctree-wrapper' in node.get('classes', None):
            numbered = node.get('numbered', 0)
            if numbered > 0 and self.numbered_level <= self.section_level:
                self.numbered = 0

    def visit_glossary(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_glossary(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_title(self, node):
        # type: (nodes.Node) -> None
        if isinstance(node.parent, nodes.topic):
            pass
        elif isinstance(node.parent, nodes.sidebar):
            pass
        elif isinstance(node.parent, nodes.Admonition):
            pass
        elif isinstance(node.parent, nodes.table):
            prefix = self._fignum_prefix(node.parent)
            self.p = self._add_paragraph(prefix + ' ', style=self.stylename['table_caption'])
            self.p.paragraph_format.keep_with_next = True
        elif isinstance(node.parent, nodes.document):
            self.docx.add_heading(node.astext().replace('\n', ' '), 0)
        elif isinstance(node.parent, nodes.section):
            headinglevel = self.section_level + self.initial_header_level - 1
            breaklevel = self.builder.config.docx_pagebreak_level
            if breaklevel is not None and headinglevel <= breaklevel:
                lastp = self.docx.paragraphs[-1] if len(self.docx.paragraphs) > 0 else None
                if lastp:
                    lastp.add_run().add_break(WD_BREAK.PAGE)
            p = self.docx.add_heading(node.astext().replace('\n', ' '), headinglevel)
            secnumlevel = self.section_level - self.numbered_level
            if self.numbered and self.numbered > secnumlevel - 1:
                self._multilevel_list_numbering(p, secnumlevel - 1, self.section_numIds[-2])
        else:
            pass

    def depart_title(self, node):
        # type: (nodes.Node) -> None
        self.p = None

    def visit_subtitle(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_subtitle(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_attribution(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_attribution(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_signature(self, node):
        # type: (nodes.Node) -> None
        self.p = self._add_paragraph()
        pass

    def depart_desc_signature(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        pass

    def visit_desc_signature_line(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc_signature_line(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_name(self, node):
        # type: (nodes.Node) -> None
        self.r = self.p.add_run()
        self.r_style = self.stylename['strong']

    def depart_desc_name(self, node):
        # type: (nodes.Node) -> None
        self.r = None
        self.r_style = None

    def visit_desc_addname(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc_addname(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_type(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc_type(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_returns(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc_returns(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_parameterlist(self, node):
        # type: (nodes.Node) -> None
        params = [child.astext() for child in node.children]
        text = '(' + ', '.join(params) + ')'
        self.r = self.p.add_run(text)

    def depart_desc_parameterlist(self, node):
        # type: (nodes.Node) -> None
        self.r = None

    def visit_desc_parameter(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_desc_optional(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc_optional(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_annotation(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_desc_annotation(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_desc_content(self, node):
        # type: (nodes.Node) -> None
        self.p_level += 1
        self.p = self._add_paragraph()
        pass

    def depart_desc_content(self, node):
        # type: (nodes.Node) -> None
        self.p_level -= 1
        self.p = None
        pass

    def visit_figure(self, node):
        # type: (nodes.Node) -> None
        self.p = self._add_paragraph()
        self.p.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.p.paragraph_format.keep_with_next = True
        self.r = self.p.add_run()

    def depart_figure(self, node):
        # type: (nodes.Node) -> None
        self.r = None
        self.p = None

    def visit_caption(self, node):
        # type: (nodes.Node) -> None
        prefix = self._fignum_prefix(node.parent)
        if isinstance(node.parent, nodes.figure):
            self.p = self._add_paragraph(prefix + u' ', style=self.stylename['image_caption'])
            self.p.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif isinstance(node.parent, nodes.container):
            self.p = self._add_paragraph(prefix + u' ', style=self.stylename['code_block_caption'])
            self.p.paragraph_format.keep_with_next = True
            pass

    def depart_caption(self, node):
        # type: (nodes.Node) -> None
        self.p = None

    def visit_productionlist(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_footnote(self, node):
        # type: (nodes.Node) -> None
        text =  node.children[0].astext().strip()
        self.p = self._add_paragraph('[%s] ' % (text))

    def depart_footnote(self, node):
        # type: (nodes.Node) -> None
        self.p = None

    def visit_citation(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_citation(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_label(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_legend(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_legend(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_option_list(self, node):
        # type: (nodes.Node) -> None
        self._add_paragraph_between_table(node)
        item_num = len(node.children)
        table = self.p_parents[-1].add_table(rows=item_num, cols=2)
        twidth = sum([cell.width for cell in table.row_cells(0)])
        for i in range(item_num):
            table.column_cells(0)[i].width = int(twidth * (1 - self.item_width_rate))
            table.column_cells(1)[i].width = int(twidth * (self.item_width_rate))
        self.tables.append([table, 0, 0])

    def depart_option_list(self, node):
        # type: (nodes.Node) -> None
        self.tables.pop()

    def visit_option_list_item(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_option_list_item(self, node):
        # type: (nodes.Node) -> None
        self.tables[-1][1] += 1
        self.tables[-1][2] = 0

    def visit_option_group(self, node):
        # type: (nodes.Node) -> None
        row = self.tables[-1][1]
        col = self.tables[-1][2]
        table = self.tables[-1][0]
        cell = table.cell(row, col)
        self.p_parents.append(cell)
        self.p = cell.paragraphs[0]

    def depart_option_group(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_parents.pop()
        self.tables[-1][2] += 1

    def visit_option(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_option(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_option_string(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_option_string(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_option_argument(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_option_argument(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_description(self, node):
        # type: (nodes.Node) -> None
        row = self.tables[-1][1]
        col = self.tables[-1][2]
        table = self.tables[-1][0]
        cell = table.cell(row, col)
        self.p_parents.append(cell)
        self.p = cell.paragraphs[0]

    def depart_description(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_parents.pop()
        self.tables[-1][2] += 1

    def visit_tabular_col_spec(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_colspec(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_tgroup(self, node):
        # type: (nodes.Node) -> None
        tgroup_node = node
        thead_num = len(tgroup_node.children[-2].children)
        tbody_num = len(tgroup_node.children[-1].children)
        row_num = thead_num + tbody_num
        col_num = tgroup_node['cols']
        table = self.p_parents[-1].add_table(rows=row_num, cols=col_num)
        align = tgroup_node.parent.get('align')
        if not align:
            align = self.builder.config.docx_imagetable_align
        if align:
            if align == 'left':
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
            elif align == 'center':
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
            elif align == 'right':
                table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        for raw in table.rows:
            for c in raw._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0
        if thead_num == 0:
            table.style = self.stylename['table'][0]
        else:
            table.style = self.stylename['table'][1]
            for i in range(thead_num):
                trPr = table.rows[i]._tr.get_or_add_trPr()
                tblHeader = trPr.get_or_add_tblHeader()
        self.tables.append([table, 0, 0])

    def depart_tgroup(self, node):
        # type: (nodes.Node) -> None
        self.tables.pop()

    def visit_thead(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_thead(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_tbody(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_tbody(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_row(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_row(self, node):
        # type: (nodes.Node) -> None
        self.tables[-1][1] += 1
        self.tables[-1][2] = 0

    def visit_entry(self, node):
        # type: (nodes.Node) -> None
        row = self.tables[-1][1]
        col = self.tables[-1][2]
        table = self.tables[-1][0]
        cell = table.cell(row, col)
        self.p_parents.append(cell)
        self.p = cell.paragraphs[0]
        if 'morerows' in node or 'morecols' in node:
            mrow = node.get('morerows', 0)
            mcol = node.get('morecols', 0)
            cell.merge(table.cell(row + mrow, col + mcol))

    def depart_entry(self, node):
        # type: (nodes.Node) -> None
        self.p_parents.pop()
        self.tables[-1][2] += 1
        col = self.tables[-1][2]
        row = self.tables[-1][1]
        table = self.tables[-1][0]
        if row > 0:
            for x in range(col + 1, len(table.rows[row].cells) - 1):
                head = table.cell(row - 1, x)
                now = table.cell(row, x)
                if head._tc != now._tc:
                    break
                self.tables[-1][2] += 1
        col = self.tables[-1][2]
        for x in range(col, len(table.rows[row].cells) - 1):
            next1 = table.cell(row, x - 1)
            next2 = table.cell(row, x)
            if next1._tc != next2._tc:
                break
            self.tables[-1][2] += 1

    def _add_paragraph_between_table(self, node):
        index = node.parent.index(node)
        prev_node = node.parent[index - 1]
        if (isinstance(prev_node, nodes.table)
            or isinstance(prev_node, nodes.field_list)
            or isinstance(prev_node, nodes.option_list)):
            self.docx.add_paragraph('')

    def visit_table(self, node):
        # type: (nodes.Node) -> None
        self._add_paragraph_between_table(node)

    def depart_table(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_acks(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_image(self, node):
        # type: (nodes.Node) -> None
        atts = {}
        uri = node['uri']
        ext = os.path.splitext(uri)[1].lower()
        if 'width' in node:
            atts['width'] = node['width']
        if 'height' in node:
            atts['height'] = node['height']
        if 'scale' in node:
            pass
        image_fullpath = os.path.join(self.builder.srcdir, uri)
        block_width = self.docx._block_width
        if isinstance(node.parent, nodes.substitution_definition):
            pass
        else:
            if isinstance(node.parent, nodes.paragraph):
                pic = self.r.add_picture(image_fullpath)
            elif isinstance(node.parent, nodes.figure):
                pic = self.r.add_picture(image_fullpath)
            else:
                p = self._add_paragraph()
                r = p.add_run()
                pic = r.add_picture(image_fullpath)
                align = node.get('align')
                if not align:
                    align = self.builder.config.docx_imagetable_align
                if align:
                    if align == 'left':
                        p.alignment = WD_TABLE_ALIGNMENT.LEFT
                    elif align == 'center':
                        p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    elif align == 'right':
                        p.alignment = WD_TABLE_ALIGNMENT.RIGHT
            if pic.width > block_width:
                pic.height = int(pic.height * float(block_width) / pic.width)
                pic.width = block_width
        raise nodes.SkipNode

    def visit_transition(self, node):
        # type: (nodes.Node) -> None
        # TODO: change from style to image
        self.docx.add_paragraph('', style=self.stylename['transition'])
        self.p = None
        raise nodes.SkipNode

    def visit_bullet_list(self, node):
        # type: (nodes.Node) -> None
        self.p_level += 1
        numId = self._get_new_num(abstractNumId=11)
        self.numIds.append(numId)

    def depart_bullet_list(self, node):
        # type: (nodes.Node) -> None
        self.p_level -= 1
        self.numIds.pop()

    def visit_enumerated_list(self, node):
        # type: (nodes.Node) -> None
        self.p_level += 1
        numId = self._get_new_num(abstractNumId=15)
        self.numIds.append(numId)

    def depart_enumerated_list(self, node):
        # type: (nodes.Node) -> None
        self.p_level -= 1
        self.numIds.pop()

    def visit_definition_list(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_definition_list(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_list_item(self, node):
        # type: (nodes.Node) -> None
        self.is_first_list_item = True

    def depart_list_item(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_definition_list_item(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_definition_list_item(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_term(self, node):
        # type: (nodes.Node) -> None
        self.p_style.append(self.stylename['definition_list'])
        self.p = self._add_paragraph(style=self.p_style[-1])

    def depart_term(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_style.pop()

    def visit_classifier(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_classifier(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_definition(self, node):
        # type: (nodes.Node) -> None
        self.p_level += 1

    def depart_definition(self, node):
        # type: (nodes.Node) -> None
        self.p_level -= 1

    def visit_field_list(self, node):
        # type: (nodes.Node) -> None
        self._add_paragraph_between_table(node)
        item_num = len(node.children)
        table = self.p_parents[-1].add_table(rows=item_num, cols=2)
        twidth = sum([cell.width for cell in table.row_cells(0)])
        for i in range(item_num):
            table.column_cells(0)[i].width = int(twidth * (1 - self.item_width_rate))
            table.column_cells(1)[i].width = int(twidth * (self.item_width_rate))
        self.tables.append([table, 0, 0])

    def depart_field_list(self, node):
        # type: (nodes.Node) -> None
        self.tables.pop()

    def visit_field(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_field(self, node):
        # type: (nodes.Node) -> None
        self.tables[-1][1] += 1
        self.tables[-1][2] = 0

    def visit_field_name(self, node):
        # type: (nodes.Node) -> None
        row = self.tables[-1][1]
        col = self.tables[-1][2]
        table = self.tables[-1][0]
        cell = table.cell(row, col)
        self.p_parents.append(cell)
        self.p = cell.paragraphs[0]

    def depart_field_name(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_parents.pop()
        self.tables[-1][2] += 1

    def visit_field_body(self, node):
        # type: (nodes.Node) -> None
        row = self.tables[-1][1]
        col = self.tables[-1][2]
        table = self.tables[-1][0]
        cell = table.cell(row, col)
        self.p_parents.append(cell)
        self.p = cell.paragraphs[0]

    def depart_field_body(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_parents.pop()
        self.tables[-1][2] += 1

    def visit_centered(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_centered(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_hlist(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_hlist(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_hlistcol(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_hlistcol(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_admonition(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_admonition(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_versionmodified(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_versionmodified(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_literal_block(self, node):
        # type: (nodes.Node) -> None
        self.p_style.append(self.stylename['literal_block'])
        self.p = self._add_paragraph(style=self.p_style[-1])

    def depart_literal_block(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_style.pop()

    def visit_doctest_block(self, node):
        # type: (nodes.Node) -> None
        self.p_style.append(self.stylename['doctest_block'])
        self.p = self._add_paragraph(style=self.p_style)

    def depart_doctest_block(self, node):
        # type: (nodes.Node) -> None
        self.p = None
        self.p_style.pop()

    def visit_line_block(self, node):
        # type: (nodes.Node) -> None
        self.p_style.append(self.stylename['line_block'])

    def depart_line_block(self, node):
        # type: (nodes.Node) -> None
        self.p_style.pop()

    def visit_line(self, node):
        # type: (nodes.Node) -> None
        text = node.astext()
        self.p = self._add_paragraph(text, style=self.p_style)

    def depart_line(self, node):
        # type: (nodes.Node) -> None
        self.p = None

    def visit_block_quote(self, node):
        # type: (nodes.Node) -> None
        self.p_style.append(self.stylename['block_quote'])
        self.p_level += 1

    def depart_block_quote(self, node):
        # type: (nodes.Node) -> None
        self.p_style.pop()
        self.p_level -= 1

    def visit_compact_paragraph(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_compact_paragraph(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_paragraph(self, node):
        # type: (nodes.Node) -> None
        if not self.p:
            self.p = self._add_paragraph(style=self.p_style)
        if isinstance(node.parent, nodes.list_item):
            if self.is_first_list_item:
                self._multilevel_list_numbering(self.p, self.p_level - 1, self.numIds[-1])
            else:
                self._multilevel_list_numbering(self.p, self.p_level - 1, 15)
            self.is_first_list_item = False
        self.r = self.p.add_run()

    def depart_paragraph(self, node):
        # type: (nodes.Node) -> None
        self.r = None
        self.p = None

    def visit_target(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_index(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_toctree(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_substitution_definition(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_pending_xref(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_pending_xref(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_reference(self, node):
        # type: (nodes.Node) -> None
        # TODO: add hyperlink
        self.r_style = self.stylename['reference']
        self.r = self.p.add_run(style=self.r_style)

    def depart_reference(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_number_reference(self, node):
        # type: (nodes.Node) -> None
        text = node.children[0].astext()
        self.r = self.p.add_run(text, style=self.r_style)
        raise nodes.SkipNode

    def visit_download_reference(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_download_reference(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_emphasis(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['emphasis']

    def depart_emphasis(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_literal_emphasis(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['literal_emphasis']

    def depart_literal_emphasis(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_strong(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['strong']

    def depart_strong(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_literal_strong(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_literal_strong(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_abbreviation(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_abbreviation(self, node):
        # type: (nodes.Node) -> None
        if node.hasattr('explanation'):
            pass

    def visit_manpage(self, node):
        # type: (nodes.Node) -> Any
        return self.visit_literal_emphasis(node)

    def depart_manpage(self, node):
        # type: (nodes.Node) -> Any
        return self.depart_literal_emphasis(node)

    def visit_title_reference(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['title_reference']

    def depart_title_reference(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_literal(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['literal']

    def depart_literal(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_subscript(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['subscript']

    def depart_subscript(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_superscript(self, node):
        # type: (nodes.Node) -> None
        self.r_style = self.stylename['superscript']

    def depart_superscript(self, node):
        # type: (nodes.Node) -> None
        self.r_style = None

    def visit_footnote_reference(self, node):
        # type: (nodes.Node) -> None
        text = node.astext()
        #text = nodes.Text(node.get('title', '#'))
        self.r = self.p.add_run('[%s]' % (text), style=self.stylename['footnote_reference'])
        raise nodes.SkipNode

    def visit_citation_reference(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_Text(self, node):
        # type: (nodes.Node) -> None
        if self.p:
            if isinstance(node.parent, nodes.field_name):
                text = node.astext() + ':'
            elif isinstance(node.parent, nodes.literal_block):
                text = node.astext().replace('\n\n', '\n')
            elif isinstance(node.parent, nodes.doctest_block):
                text = node.astext().replace('\n\n', '\n')
            else:
                text = node.astext().replace('\n', ' ')
            self.r = self._add_run(text, style=self.r_style)

    def depart_Text(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_generated(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_generated(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_inline(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_inline(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_container(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_container(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_problematic(self, node):
        # type: (nodes.Node) -> None
        pass

    def depart_problematic(self, node):
        # type: (nodes.Node) -> None
        pass

    def visit_system_message(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_comment(self, node):
        # type: (nodes.Node) -> None
        raise nodes.SkipNode

    def visit_meta(self, node):
        # type: (nodes.Node) -> None
        # only valid for HTML
        raise nodes.SkipNode

    def visit_raw(self, node):
        # type: (nodes.Node) -> None
        # TODO: support raw HTML translation
        raise nodes.SkipNode

    def visit_math(self, node):
        # type: (nodes.Node) -> None
        # TODO: support to convert from latex to docx
        eq = node.get('latex')
        text = eq if eq else node.astext()
        self.r = self._add_run(text, style=self.r_style)
        self.r = None
        raise nodes.SkipNode

    def visit_math_block(self, node):
        # type: (nodes.Node) -> None
        self.visit_displaymath(node)

    def depart_math_block(self, node):
        # type: (nodes.Node) -> None
        self.depart_displaymath(node)

    def visit_displaymath(self, node):
        # type: (nodes.Node) -> None
        eq = node.get('latex')
        eq = eq if eq else node.astext()
        number = node.get('number')
        number = '(%s)' % str(number) if number else ''
        self._add_paragraph_between_table(node)
        table = self.p_parents[-1].add_table(rows=1, cols=3)
        twidth = sum([cell.width for cell in table.row_cells(0)])
        table.cell(0, 0).width = int(twidth * 0.1)
        table.cell(0, 0).text = ''
        table.cell(0, 0).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.LEFT
        table.cell(0, 1).width = int(twidth * 0.8)
        table.cell(0, 1).text = eq
        table.cell(0, 1).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 2).width = int(twidth * 0.1)
        table.cell(0, 2).text = number
        table.cell(0, 2).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.RIGHT

    def depart_displaymath(self, node):
        # type: (nodes.Node) -> None
        self.p = None

    def unknown_visit(self, node):
        # type: (nodes.Node) -> None
        raise NotImplementedError('Unknown node: ' + node.__class__.__name__)
